
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\DU-AN-PHAN-MEM\PHUBAI-MES\phubai-mes\PHUBAI-MES-Quy-trinh-xay-dung-va-deploy.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
TEXT = "1F2937"
MUTED = "4B5563"
LIGHT = "E8EEF5"
BORDER = "B7C7D9"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def borders(table):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in('w:tblBorders')
    if b is None:
        b = OxmlElement('w:tblBorders')
        tbl_pr.append(b)
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        e = b.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge)
            b.append(e)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), BORDER)


def table_width(table):
    tbl_pr = table._tbl.tblPr
    w = tbl_pr.first_child_found_in('w:tblW')
    if w is None:
        w = OxmlElement('w:tblW')
        tbl_pr.append(w)
    w.set(qn('w:w'), '9360')
    w.set(qn('w:type'), 'dxa')
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')


def cell_margins(table):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.first_child_found_in('w:tblCellMar')
    if mar is None:
        mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(mar)
    for name, value in [('top', 80), ('bottom', 80), ('start', 120), ('end', 120)]:
        node = mar.find(qn('w:' + name))
        if node is None:
            node = OxmlElement('w:' + name)
            mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_widths(table, widths):
    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width * 1440)))
            tc_w.set(qn('w:type'), 'dxa')


def set_cell(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Calibri'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    table_width(t)
    set_widths(t, widths)
    cell_margins(t)
    borders(t)
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], LIGHT)
        set_cell(t.rows[0].cells[i], h, True, DARK)
    for values in rows:
        row = t.add_row()
        for i, v in enumerate(values):
            set_cell(row.cells[i], v)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def p(doc, text='', style=None):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def nums(doc, items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def code(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.15)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    pr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F3F4F6')
    pr.append(shd)
    lines = text.strip('\n').split('\n')
    for i, line in enumerate(lines):
        if i:
            para.add_run('\n')
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string('111827')


def note(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width(t)
    cell_margins(t)
    borders(t)
    c = t.cell(0, 0)
    shade(c, 'F4F6F9')
    c.text = ''
    para = c.paragraphs[0]
    r = para.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    r.font.size = Pt(10.5)
    para2 = c.add_paragraph(body)
    para2.paragraph_format.space_after = Pt(0)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [('Heading 1', 16, BLUE, 18, 10), ('Heading 2', 13, BLUE, 14, 7), ('Heading 3', 12, DARK, 10, 5)]:
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    setup_styles(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('PHUBAI-MES')
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(DARK)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('Quy trình xây dựng phần mềm từ đầu và deploy tự động bằng GitHub Actions').font.size = Pt(13)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run('Workspace: D:\\DU-AN-PHAN-MEM\\PHUBAI-MES\\phubai-mes | Port: 3002 | Database: phubai_mes_db')
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = RGBColor.from_string(MUTED)
    note(doc, 'Mục tiêu tài liệu', 'Tài liệu này tóm tắt toàn bộ quy trình dựng PHUBAI-MES từ đầu: chuẩn bị stack, tạo project, cấu hình database, đưa code lên GitHub, deploy tự động bằng GitHub Actions, chạy PM2 và kết nối domain Cloudflare Tunnel về localhost:3002.')

    doc.add_heading('1. Kiến trúc tổng quan', level=1)
    p(doc, 'PHUBAI-MES là ứng dụng MES mới, độc lập với PHUBAI-ERP. Module đầu tiên là điện năng, kế thừa nghiệp vụ ERP nhưng dùng database, namespace và deploy riêng.')
    add_table(doc, ['Thành phần', 'Giá trị chuẩn'], [
        ['Framework', 'Next.js App Router, TypeScript, React'],
        ['UI', 'Ant Design'],
        ['Database', 'PostgreSQL + Prisma ORM'],
        ['Auth', 'NextAuth.js v5 / Auth.js'],
        ['Runtime', 'Node.js 22 trên Windows Server'],
        ['Process manager', 'PM2: phubai-mes-web và phubai-mes-energy-cron'],
        ['Port', '3002'],
        ['Domain', 'https://phubaimes.site qua Cloudflare Tunnel'],
        ['Namespace điện năng', '/electric và /api/electric/*'],
    ], [1.7, 4.8])

    doc.add_heading('2. Chuẩn bị môi trường phát triển', level=1)
    nums(doc, ['Cài Node.js LTS, khuyến nghị Node 22.', 'Cài Git và cấu hình tài khoản GitHub.', 'Cài PostgreSQL và pgAdmin nếu cần.', 'Tạo workspace D:\\DU-AN-PHAN-MEM\\PHUBAI-MES.', 'Chuẩn bị tài khoản Cloudflare và GitHub repository.'])
    code(doc, 'node -v\ngit --version\npsql --version')

    doc.add_heading('3. Tạo folder và khởi tạo project', level=1)
    code(doc, 'cd D:\\DU-AN-PHAN-MEM\nmkdir PHUBAI-MES\ncd D:\\DU-AN-PHAN-MEM\\PHUBAI-MES\nnpx create-next-app@latest phubai-mes --ts --app\ncd D:\\DU-AN-PHAN-MEM\\PHUBAI-MES\\phubai-mes')
    p(doc, 'Cố định port 3002 trong package.json để MES chạy song song ERP/HRM.')
    code(doc, '"dev": "next dev -p 3002"\n"build": "next build"\n"start": "next start -p 3002"')

    doc.add_heading('4. Cài stack chính', level=1)
    code(doc, 'npm install antd @ant-design/icons next-auth@beta @prisma/client @prisma/adapter-pg pg bcryptjs dayjs node-cron modbus-serial xlsx pm2\nnpm install -D prisma typescript eslint eslint-config-next tailwindcss @tailwindcss/postcss tsx\nnpx prisma init')
    note(doc, 'Lưu ý build trên GitHub Actions', 'Không đặt NODE_ENV=production trước npm ci, vì Next build vẫn cần devDependencies như @tailwindcss/postcss. Workflow dùng npm ci --include=dev.')

    doc.add_heading('5. Database và .env', level=1)
    code(doc, 'createdb -U postgres phubai_mes_db')
    code(doc, 'DATABASE_URL="postgresql://postgres:123456@localhost:5432/phubai_mes_db?schema=public"\nNEXTAUTH_URL="http://localhost:3002"\nNEXTAUTH_SECRET="TAO_CHUOI_SECRET_DAI_TOI_THIEU_32_KY_TU"\nAUTH_SECRET="TAO_CHUOI_SECRET_DAI_TOI_THIEU_32_KY_TU"\nAUTH_TRUST_HOST=true')

    doc.add_heading('6. Schema và module điện năng', level=1)
    p(doc, 'Module điện năng MES dùng cấu trúc Factory -> PowerTransformer -> PowerMeter để thống kê điện năng/tiền điện theo từng nhà máy.')
    add_table(doc, ['Model', 'Vai trò'], [
        ['Factory', 'Danh mục nhà máy.'], ['PowerTransformer', 'Trạm biến áp, thuộc nhà máy.'], ['PowerMeterGroup', 'Nhóm đồng hồ.'], ['PowerMeter', 'Đồng hồ điện, AUTO/MANUAL, Gateway IP, Modbus ID, TU/TI.'], ['ElectricityPrice', 'Giá điện, mặc định NORMAL.'], ['PowerTelemetry', 'Dữ liệu thô AUTO/realtime.'], ['PowerRecord', 'Dữ liệu chốt ngày, AUTO hoặc MANUAL.']
    ], [1.7, 4.8])
    code(doc, 'npx prisma generate\nnpx prisma migrate dev --name init_energy_module\nnpx prisma migrate dev --name add_factory_electric_hierarchy')
    bullets(doc, ['UI chính: /electric/overview, /electric/catalog, /electric/daily-input, /electric/live, /electric/reports, /electric/prices.', 'API chính: /api/electric/*.', 'Không gộp toàn bộ module vào một trang; dùng layout/sidebar riêng cho điện năng.'])

    doc.add_heading('7. Cron điện năng và kiểm tra AUTO', level=1)
    p(doc, 'Cron điện năng dùng scripts/energy-cron.js, chạy bằng Node JS thuần. Logic đọc AUTO giống ERP: nhóm đồng hồ theo Gateway, mở TCP một lần cho mỗi Gateway, đọc từng Modbus ID bằng client.setID(), nghỉ 50ms giữa mỗi lần đọc.')
    code(doc, 'npm run energy:cron -- --status\nnpm run energy:cron -- --collect-once\nnpm run energy:cron -- --close-once')
    add_table(doc, ['Kết quả', 'Ý nghĩa'], [['AUTO meters > 0', 'Đã có đồng hồ AUTO.'], ['inserted > 0', 'Đọc được đồng hồ và ghi PowerTelemetry.'], ['Lỗi Gateway', 'Sai IP/port hoặc server chưa vào được USR-N520.'], ['THIEU_CAU_HINH', 'Thiếu gatewayIp hoặc modbusId.']], [1.9, 4.6])

    doc.add_heading('8. Đưa code lên GitHub', level=1)
    code(doc, 'git init\ngit add .\ngit commit -m "Initialize PHUBAI-MES"\ngit branch -M main\ngit remote add origin https://github.com/<user>/phubai-mes.git\ngit push -u origin main')
    p(doc, 'Quy trình update hằng ngày: sửa code, chạy kiểm tra, commit, push main để GitHub Actions tự deploy.')
    code(doc, 'npm run lint\nnpm run build\ngit status\ngit add .\ngit commit -m "Mo ta thay doi"\ngit push origin main')

    doc.add_heading('9. GitHub Actions self-hosted runner', level=1)
    nums(doc, ['Vào GitHub repo -> Settings -> Actions -> Runners -> New self-hosted runner.', 'Chọn Windows và chạy lệnh config trên server.', 'Thêm label phubai-mes.', 'Nếu PM2 lỗi quyền pipe, đổi service runner sang Windows user thật thay vì NetworkService.', 'Runner phải hiện Idle trước khi deploy.'])

    doc.add_heading('10. Repository Secrets', level=1)
    p(doc, 'Tạo từng Repository Secret riêng, không gộp nhiều dòng vào một secret.')
    add_table(doc, ['Secret', 'Giá trị mẫu'], [['DATABASE_URL', 'postgresql://postgres:123456@localhost:5432/phubai_mes_db?schema=public'], ['NEXTAUTH_URL', 'https://phubaimes.site'], ['NEXTAUTH_SECRET', 'Chuỗi bí mật tối thiểu 32 ký tự'], ['AUTH_SECRET', 'Có thể dùng cùng NEXTAUTH_SECRET'], ['AUTH_TRUST_HOST', 'true']], [1.8, 4.7])
    code(doc, 'node -e "console.log(require(\'crypto\').randomBytes(32).toString(\'base64\'))"')

    doc.add_heading('11. Workflow deploy tự động', level=1)
    p(doc, 'Workflow .github/workflows/deploy.yml được thiết kế giống ERP: trước checkout chỉ dừng app MES, sau build chỉ start/reload app MES bằng --only, không động ERP/HRM.')
    add_table(doc, ['Bước', 'Nội dung'], [['1', 'Stop phubai-mes-web và phubai-mes-energy-cron; chỉ giải phóng port 3002.'], ['2', 'Checkout source.'], ['3', 'Setup Node 22.'], ['4', 'Validate secrets.'], ['5', 'Tạo .env từ secrets.'], ['6', 'npm ci --include=dev.'], ['7', 'npx prisma generate.'], ['8', 'npx prisma migrate deploy.'], ['9', 'npm run build.'], ['10', 'startOrReload chỉ phubai-mes-web và phubai-mes-energy-cron.'], ['11', 'Fallback cũng chỉ restart MES apps.']], [0.8, 5.7])
    code(doc, 'pm2 startOrReload ecosystem.config.cjs --only phubai-mes-web --update-env\npm2 startOrReload ecosystem.config.cjs --only phubai-mes-energy-cron --update-env')

    doc.add_heading('12. PM2 ecosystem', level=1)
    add_table(doc, ['PM2 app', 'Chức năng'], [['phubai-mes-web', 'Next.js production: next start -p 3002.'], ['phubai-mes-energy-cron', 'Node chạy scripts/energy-cron.js để thu telemetry và chốt số.']], [2.0, 4.5])
    bullets(doc, ['Dùng PM2_HOME trong workspace để không đụng PM2 của ERP/HRM.', 'ecosystem.config.cjs dùng __dirname làm cwd.', 'windowsHide: true để không bật cửa sổ node.exe trên Windows.'])

    doc.add_heading('13. Cloudflare Tunnel và domain', level=1)
    nums(doc, ['Đưa phubaimes.site vào Cloudflare và đổi nameserver tại nơi mua domain.', 'Vào Zero Trust -> Networks -> Tunnels.', 'Mở tunnel server-phubai nếu ERP/HRM đã chạy cùng server; không cần tạo tunnel mới.', 'Add Public Hostname cho MES.', 'Subdomain để trống nếu dùng domain gốc phubaimes.site.', 'Domain chọn phubaimes.site.', 'Service type HTTP, URL localhost:3002 hoặc http://localhost:3002.', 'Nếu báo A/AAAA/CNAME tồn tại, xóa DNS record @ hoặc phubaimes.site cũ rồi thêm lại.'])
    code(doc, 'Hostname: phubaimes.site\nService: http://localhost:3002\nTest: https://phubaimes.site/electric/overview')

    doc.add_heading('14. Chuyển dữ liệu điện năng từ ERP sang MES', level=1)
    p(doc, 'Dùng script import database-to-database. Script không ghi ngược vào ERP, dùng upsert và map ID dạng erp-* để giữ quan hệ.')
    code(doc, "$env:ERP_DATABASE_URL='postgresql://postgres:123456@localhost:5432/phubai_erp_db?schema=public'\n$env:DATABASE_URL='postgresql://postgres:123456@localhost:5432/phubai_mes_db?schema=public'\nnpm run energy:import-erp -- --dry-run\nnpm run energy:import-erp -- --yes")
    note(doc, 'Bắt buộc backup', 'Trước khi import thật trên production, backup phubai_mes_db bằng pg_dump. Chỉ chạy --yes khi dry-run cho số dòng đúng.')

    doc.add_heading('15. Checklist triển khai', level=1)
    bullets(doc, ['Database phubai_mes_db đã tồn tại.', 'Runner GitHub hiển thị Idle và có label phubai-mes.', 'Secrets đã đủ.', 'Workflow deploy.yml đã push lên main.', 'PM2 apps MES đang online.', 'Cloudflare Public Hostname trỏ về http://localhost:3002.', 'Truy cập được /electric/overview.', 'Cron --status hiển thị đồng hồ AUTO nếu đã cấu hình.'])

    doc.add_heading('16. Lỗi thường gặp', level=1)
    add_table(doc, ['Lỗi', 'Nguyên nhân', 'Cách xử lý'], [['DATABASE_URL is empty', 'Thiếu Repository Secret.', 'Tạo DATABASE_URL trong Repository secrets.'], ['Cannot find @tailwindcss/postcss', 'npm ci thiếu devDependencies.', 'Dùng npm ci --include=dev.'], ['pm2 not recognized', 'Runner không thấy PM2 global.', 'Dùng pm2 dependency/local binary.'], ['connect EPERM rpc.sock', 'Runner chạy bằng NetworkService.', 'Đổi service runner sang Windows user thật.'], ['Cửa sổ node.exe bật lên', 'PM2 spawn process tương tác.', 'Thêm windowsHide: true.'], ['Cloudflare record exists', 'Có DNS record trùng hostname.', 'Xóa record cũ rồi add Public Hostname lại.'], ['Cron inserted = 0', 'Chưa có AUTO meter hoặc sai gateway.', 'Kiểm tra /electric/catalog và --status.']], [1.55, 2.25, 2.7])

    doc.add_heading('17. File nguồn cần đọc khi tiếp tục phát triển', level=1)
    bullets(doc, ['AGENTS.md', 'BUSINESS_LOGIC_CONTEXT.md', 'PLANS/yeucau.md', 'PROJECT_SKILLS/phubai-mes-electric/SKILL.md', 'prisma/schema.prisma', 'scripts/energy-cron.js', '.github/workflows/deploy.yml'])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run('PHUBAI-MES - Quy trình xây dựng và deploy')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string(MUTED)
    doc.save(OUT)

if __name__ == '__main__':
    build()
